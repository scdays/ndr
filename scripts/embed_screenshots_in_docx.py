#!/usr/bin/env python3
"""Build V1.1 Word report with embedded evidence screenshots."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

ASSETS = Path("/workspace/assets/evidence_screenshots")
SRC_DOC = Path("/workspace/NDR市场加密流量检测深度分析报告V1.1.docx")
OUT_DOC = Path("/workspace/NDR市场加密流量检测深度分析报告V1.1-含截图.docx")

# Section anchor -> list of (screenshot_id, caption, source_url)
EMBEDDINGS = [
    ("1.1 市场规模", [
        ("S-03", "图1  IDC 2022：中国 NDR 市场规模 3.5 亿美元（+13.7%）", "https://www.secrss.com/articles/58067"),
        ("S-01", "图2  IDC 2024：中国 NDR 市场规模 23.6 亿元（-0.9%）", "https://www.secrss.com/articles/79613"),
    ]),
    ("1.2 市场竞争", [
        ("S-02", "图3  IDC 2023：奇安信市场份额 22.5%（连续三年第一）", "https://www.qianxin.com/report/detail/rid/62"),
        ("S-01", "图4  IDC 2024：主要厂商及加密流量检测趋势", "https://www.secrss.com/articles/79613"),
    ]),
    ("2.1 全球加密", [
        ("S-04", "图5  Zscaler ThreatLabz 2024：87.2% 攻击经加密通道投递", "https://www.zscaler.com/blogs/security-research/threatlabz-report-threats-delivered-over-encrypted-channels"),
    ]),
    ("2.2 全球加密流量渗透率", [
        ("S-05", "图6  HTTP Archive 2024：98% Web 请求使用 HTTPS", "https://almanac.httparchive.org/en/2024/security"),
    ]),
    ("2.3 IDC", [
        ("S-01", "图7  IDC 2024：加密流量检测成为核心趋势", "https://www.secrss.com/articles/79613"),
    ]),
    ("3.1 全球 NDR", [
        ("S-06", "图8  Grand View Research：2025 年全球 NDR 市场 37.7 亿美元", "https://www.grandviewresearch.com/industry-analysis/network-detection-response-ndr-market-report"),
        ("S-07", "图9  MarketsandMarkets：2030 年全球 NDR 市场 58.2 亿美元预测", "https://www.marketsandmarkets.com/Market-Reports/network-detection-and-response-market-236524642.html"),
    ]),
    ("4.2 ETD", [
        ("S-10", "图10 数世咨询：NDR 能力指南行业需求分布", "https://www.dwcon.cn/post/3410"),
    ]),
    ("六、证据附件", []),  # appendix section - add all at end
]


def set_run(run, size=10, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_evidence_block(doc, sid: str, caption: str, url: str):
    img_path = ASSETS / f"{sid}.png"
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run(caption)
    set_run(r, 10, True, RGBColor(0x1F, 0x4E, 0x79))

    if img_path.exists():
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_p.add_run()
        run.add_picture(str(img_path), width=Inches(6.2))
    else:
        miss = doc.add_paragraph()
        mr = miss.add_run(f"[截图缺失：{sid}.png，请运行 capture_evidence_screenshots.py]")
        set_run(mr, 9, False, RGBColor(0xC0, 0x00, 0x00))

    p_url = doc.add_paragraph()
    p_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ur = p_url.add_run(f"来源：{url}")
    set_run(ur, 8, False, RGBColor(0x66, 0x66, 0x66))
    doc.add_paragraph()


def find_paragraph_index(doc, keyword: str) -> int:
    for i, para in enumerate(doc.paragraphs):
        if keyword in para.text:
            return i
    return -1


def insert_after_keyword(doc: Document, keyword: str, items: list):
    idx = find_paragraph_index(doc, keyword)
    if idx < 0:
        print(f"  Warning: anchor '{keyword}' not found, skip embed")
        return False

    # python-docx cannot insert mid-document easily; we rebuild by cloning approach:
    # Instead: append evidence sections at end grouped, OR regenerate full doc.
    return False


def build_full_document():
    """Regenerate complete V1.1 docx with screenshots inline."""
    from generate_v11_docx import build as build_base  # noqa

    # Import inline by exec if module name issue
    import importlib.util
    spec = importlib.util.spec_from_file_location(
        "g11", "/workspace/scripts/generate_v11_docx.py"
    )
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    mod.build()

    doc = Document(str(SRC_DOC))

    # Collect all unique embeddings in order
    all_items = []
    seen = set()
    for _anchor, items in EMBEDDINGS:
        for item in items:
            if item[0] not in seen:
                all_items.append(item)
                seen.add(item[0])

    # Add dedicated appendix section before disclaimer (end)
    doc.add_page_break()
    h = doc.add_heading("附录：数据证据截图", level=1)
    intro = doc.add_paragraph(
        "以下截图为公开渠道原文页面截图，截取时间为 2026 年 5 月，"
        "用于佐证本报告关键统计数据。完整链接见正文数据来源索引。"
    )
    for r in intro.runs:
        set_run(r, 10)

    fig_num = 1
    for sid, caption_template, url in all_items:
        caption = caption_template.replace("图N ", f"图{fig_num} ")
        if "图1 " not in caption and "图" in caption_template[:3]:
            # fix numbering
            pass
        caption = f"图{fig_num}  {caption_template.split('  ', 1)[-1] if '  ' in caption_template else caption_template}"
        add_evidence_block(doc, sid, caption, url)
        fig_num += 1

    doc.save(str(OUT_DOC))
    print(f"Saved: {OUT_DOC}")


def build_from_scratch_with_appendix():
    """Load existing v1.1 or build base, append screenshot appendix."""
    if not SRC_DOC.exists():
        import importlib.util
        spec = importlib.util.spec_from_file_location(
            "g11", "/workspace/scripts/generate_v11_docx.py"
        )
        mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)
        mod.build()

    doc = Document(str(SRC_DOC))

    doc.add_page_break()
    doc.add_heading("附录：数据证据截图", level=1)
    p = doc.add_paragraph(
        "以下截图为公开网页原文截图（2026年5月采集），对应正文各章节关键数据。"
        "若页面改版导致数值位置变化，以来源链接实时页面为准。"
    )

    items = [
        ("S-03", "图1  IDC 2022：中国 NDR 市场规模 3.5 亿美元，同比 +13.7%", "https://www.secrss.com/articles/58067"),
        ("S-01", "图2  IDC 2024：中国 NDR 市场规模 23.6 亿元，同比 -0.9%", "https://www.secrss.com/articles/79613"),
        ("S-02", "图3  IDC 2023：奇安信市场份额 22.5%，连续三年第一", "https://www.qianxin.com/report/detail/rid/62"),
        ("S-04", "图4  Zscaler ThreatLabz 2024：87.2% 威胁经加密通道投递", "https://www.zscaler.com/blogs/security-research/threatlabz-report-threats-delivered-over-encrypted-channels"),
        ("S-05", "图5  HTTP Archive 2024：98% Web 请求使用 HTTPS", "https://almanac.httparchive.org/en/2024/security"),
        ("S-06", "图6  Grand View Research：2025 年全球 NDR 37.7 亿美元", "https://www.grandviewresearch.com/industry-analysis/network-detection-response-ndr-market-report"),
        ("S-07", "图7  MarketsandMarkets：2030 年全球 NDR 58.2 亿美元（CAGR 9.6%）", "https://www.marketsandmarkets.com/Market-Reports/network-detection-and-response-market-236524642.html"),
        ("S-10", "图8  数世咨询：NDR 能力指南—行业需求分布", "https://www.dwcon.cn/post/3410"),
    ]

    for sid, caption, url in items:
        add_evidence_block(doc, sid, caption, url)

    # Also add inline callouts in summary - add note after title
    doc.save(str(OUT_DOC))
    print(f"Saved: {OUT_DOC} ({len(items)} figures)")


if __name__ == "__main__":
    build_from_scratch_with_appendix()
