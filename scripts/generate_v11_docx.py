#!/usr/bin/env python3
"""Generate NDR V1.1 report as Word document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn


def set_cell_shading(cell, color: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_table(doc, headers, rows, header_color="D9E2F3"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for r in p.runs:
        r.font.size = Pt(10)


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_heading("NDR 市场加密流量方向深度分析报告 V1.1", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("（2022–2024 历史数据 + 2026–2030 市场预测 · 证据标注版）")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run("报告版本：V1.1  |  报告日期：2026年5月  |  适用：NDR 2.0 加密流量分析需求验证").font.size = Pt(9)

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(
        "本报告聚焦全球及中国 NDR（Network Detection & Response）市场中加密流量检测（ETD）"
        "细分方向。V1.1 对 V1.0 全部关键数据逐项公开验证，附来源链接与证据等级。"
    )

    doc.add_heading("核心发现（均已公开验证）", level=2)
    add_table(doc,
        ["维度", "关键数据", "证据等级", "来源"],
        [
            ["中国 NDR 市场", "2024年 23.6亿元，-0.9%", "✅ 已验证", "IDC CHC52196225"],
            ["全球 NDR 市场", "2025年 37.7亿美元", "✅ 已验证", "Grand View Research"],
            ["全球 NDR 预测", "2030年 58.2亿美元，CAGR 9.6%", "✅ 已验证", "MarketsandMarkets"],
            ["加密流量威胁", "87.2% 经 TLS/SSL 投递", "✅ 已验证", "Zscaler ThreatLabz 2024"],
            ["恶意软件占比", "86.5%", "✅ 已验证", "Zscaler ThreatLabz 2024"],
            ["HTTPS 普及", "98% Web 请求", "✅ 已验证", "HTTP Archive 2024"],
            ["TLS 1.3", "73% 站点", "✅ 已验证", "HTTP Archive 2024"],
            ["技术趋势", "加密流量检测为核心", "✅ 已验证", "IDC CHC52196225"],
        ])

    add_note(doc,
        "V1.1 删除项：CR5=56.5%、ETD CAGR=32%、金融渗透率68%、2027年60亿元等无公开出处数据已删除。")

    # Section 1
    doc.add_heading("一、中国 NDR 市场规模与格局（2022–2024）", level=1)
    doc.add_heading("1.1 市场规模变化趋势", level=2)
    add_table(doc,
        ["年份", "市场规模", "同比增速", "来源", "证据链接"],
        [
            ["2022", "3.5亿美元（约24.2亿元）", "+13.7%", "IDC CHC50358323", "secrss.com/articles/58067"],
            ["2023", "23.8亿元人民币", "+1.6%", "IDC CHC50964624", "qianxin.com/report/detail/rid/62"],
            ["2024", "23.6亿元人民币", "-0.9%", "IDC CHC52196225", "secrss.com/articles/79613"],
        ])

    doc.add_heading("【数据冲突标注】", level=3)
    add_table(doc,
        ["机构", "2022年规模", "说明"],
        [
            ["IDC（采用）", "3.5亿美元", "出货量统计"],
            ["华经产业研究院", "25.48亿元", "厂商调研，约高5%"],
            ["数世咨询", "24.64亿元", "18家厂商调研"],
        ])

    doc.add_heading("1.2 市场竞争格局（2024）", level=2)
    add_table(doc,
        ["排名", "厂商", "公开信息"],
        [
            ["1", "奇安信", "连续四年市场份额第一"],
            ["2", "绿盟科技", "TOP5 主要玩家"],
            ["3", "深信服科技", "TOP5 主要玩家"],
            ["4", "安恒信息", "TOP5（与深信服并列第三）"],
            ["5", "360数字安全集团", "TOP5 主要玩家"],
            ["—", "观成科技", "IDC 加密流量检测代表厂商"],
        ])

    add_note(doc,
        "修正：22.5% 为 2023 年奇安信份额（非 2024）。2024 精确份额需 IDC 报告截图（S-09）。")

    # Section 2
    doc.add_heading("二、加密流量检测（ETD）专项分析", level=1)
    doc.add_heading("2.1 全球加密流量威胁态势", level=2)
    doc.add_paragraph("来源：Zscaler ThreatLabz 2024 Encrypted Attacks Report（321亿样本，2023.10–2024.9）")
    add_table(doc,
        ["指标", "数值", "证据链接"],
        [
            ["加密通道攻击占比", "87.2%（+10.3% YoY）", "zscaler.com/blogs/security-research/..."],
            ["恶意软件占比", "86.5%", "同上"],
            ["XSS 增幅", "+110.2% YoY", "同上"],
            ["制造业攻击占比", "42%", "globenewswire.com/..."],
        ])

    doc.add_heading("2.2 全球加密流量渗透率", level=2)
    add_table(doc,
        ["指标", "数值", "来源"],
        [
            ["HTTPS 请求占比", "98%", "HTTP Archive Web Almanac 2024"],
            ["HTTPS 站点占比", "96%", "HTTP Archive 2024"],
            ["TLS 1.3 站点占比", "73%", "HTTP Archive 2024"],
            ["Google 流量加密率", "~96%", "Google Transparency Report"],
        ])

    doc.add_heading("2.3 IDC 加密流量检测趋势判断", level=2)
    p = doc.add_paragraph()
    p.add_run("2024年 IDC 原文：").bold = True
    p.add_run("「加密流量检测成为更多行业的明确需求，利用 AI 实现不解密状态下的流量威胁检测"
              "成为技术提供商核心竞争力之一。」")
    doc.add_paragraph("证据：https://www.secrss.com/articles/79613")

    doc.add_heading("2.4 三大技术路线", level=2)
    add_table(doc,
        ["技术路线", "原理", "代表厂商"],
        [
            ["AI 不解密", "TLS元数据+JA3+行为+AI/ML", "观成、360、绿盟"],
            ["可选解密", "可疑流量选择性解密", "奇安信、深信服"],
            ["全量解密", "中间人全量解密", "传统安全厂商"],
        ])

    # Section 3
    doc.add_heading("三、全球 NDR 市场对比", level=1)
    add_table(doc,
        ["年份", "规模", "CAGR", "来源"],
        [
            ["2024", "33.4亿美元", "—", "MarketsandMarkets"],
            ["2025", "37.7亿美元", "—", "Grand View Research"],
            ["2030（预测）", "58.2亿美元", "9.6%", "MarketsandMarkets"],
            ["2033（预测）", "80.8亿美元", "10.1%", "Grand View Research"],
        ])
    add_note(doc, "数据冲突：Grand View（2033）与 MarketsandMarkets（2030）不可合并为单一曲线。")

    # Section 4
    doc.add_heading("四、市场预测（2026–2030）", level=1)
    doc.add_heading("4.1 中国 NDR 预测【估算】", level=2)
    add_table(doc,
        ["年份", "规模（亿元）", "增速", "依据"],
        [
            ["2026", "31.5", "+18.6%", "贝哲斯咨询"],
            ["2027", "37.5", "+19.0%", "贝哲斯咨询"],
            ["2028", "44.5", "+18.7%", "【估算】"],
            ["2029", "53.0", "+19.1%", "【估算】"],
            ["2030", "63.0", "+18.9%", "【估算】"],
        ])
    doc.add_paragraph("来源：https://m.gelonghui.com/p/901401")

    doc.add_heading("4.2 ETD 细分市场", level=2)
    doc.add_paragraph(
        "V1.1 修正：ETD CAGR=32% 等精确数字无公开出处已删除。"
        "IDC 定性判断：加密流量检测成为更多行业明确需求。"
    )

    # Section 5
    doc.add_heading("五、NDR 2.0 需求验证结论", level=1)
    conclusions = [
        "市场体量支撑：中国 NDR 23.6 亿（2024），全球 37.7 亿美元（2025）",
        "威胁驱动明确：87.2% 攻击经加密通道，86.5% 为恶意软件",
        "技术方向清晰：IDC 确认 AI 不解密检测为核心竞争力",
        "竞品窗口存在：观成/360/绿盟/奇安信/深信服均在布局",
        "待补充：竞品 PoC 实测、IDC 精确份额截图、ETD 细分量化",
    ]
    for c in conclusions:
        add_bullet(doc, c)

    # Section 6 - Screenshot list
    doc.add_heading("六、截图采集清单", level=1)
    add_table(doc,
        ["编号", "内容", "URL"],
        [
            ["S-01", "IDC 2024 规模+份额图", "secrss.com/articles/79613"],
            ["S-02", "奇安信 22.5%（2023）", "qianxin.com/report/detail/rid/62"],
            ["S-03", "IDC 2022 3.5亿美元", "secrss.com/articles/58067"],
            ["S-04", "Zscaler 87.2%", "zscaler.com/blogs/security-research/..."],
            ["S-05", "HTTP Archive 98% HTTPS", "almanac.httparchive.org/en/2024/security"],
            ["S-06", "Grand View 37.7亿", "grandviewresearch.com/..."],
            ["S-07", "M&M 2030 预测", "marketsandmarkets.com/..."],
            ["S-08", "IDC 加密流量趋势", "secrss.com/articles/79613"],
            ["S-09", "IDC 2024 份额表（付费）", "IDC CHC52196225"],
            ["S-10", "数世咨询行业分布", "dwcon.cn/post/3410"],
        ])

    # Section 7 - Sources
    doc.add_heading("七、数据来源索引", level=1)
    add_table(doc,
        ["#", "机构", "报告", "链接"],
        [
            ["1", "IDC", "NDR 市场份额 2022", "secrss.com/articles/58067"],
            ["2", "IDC", "NDR 市场份额 2023", "qianxin.com/report/detail/rid/62"],
            ["3", "IDC", "GenAI NDR 2024", "secrss.com/articles/79613"],
            ["4", "Zscaler", "Encrypted Attacks 2024", "zscaler.com/blogs/..."],
            ["5", "HTTP Archive", "Web Almanac 2024", "almanac.httparchive.org/..."],
            ["6", "Grand View", "NDR Market", "grandviewresearch.com/..."],
            ["7", "MarketsandMarkets", "NDR Market", "marketsandmarkets.com/..."],
            ["8", "数世咨询", "NDR 能力指南", "dwcon.cn/post/3410"],
            ["9", "贝哲斯", "NDR 市场规模", "m.gelonghui.com/p/901401"],
        ])

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    d = disclaimer.add_run(
        "免责声明：本报告 V1.1 数据均经公开渠道交叉验证。【估算】数据已标注依据。"
        "精确份额需 IDC 付费报告截图。报告仅供参考，不构成投资建议。"
    )
    d.font.size = Pt(8)
    d.italic = True

    out = "/workspace/NDR市场加密流量检测深度分析报告V1.1.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
