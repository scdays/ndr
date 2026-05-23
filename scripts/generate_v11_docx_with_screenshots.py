#!/usr/bin/env python3
"""Generate full NDR V1.2 Word report with screenshots and expanded chapters 4–6."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn

from report_chapters_forecast_trends import CHAPTER_4, CHAPTER_5, CHAPTER_6

ASSETS = Path("/workspace/assets/evidence_screenshots")
OUT = Path("/workspace/NDR市场加密流量检测深度分析报告V1.1-含截图.docx")
OUT_MAIN = Path("/workspace/NDR市场加密流量检测深度分析报告V1.1.docx")


def set_cell_shading(cell, color: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
    shading.append(shd)


def style_run(run, size=10, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_table(doc, headers, rows, header_color="D9E2F3"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.cell(ri + 1, ci).text = str(val)
            for p in table.cell(ri + 1, ci).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, 9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    style_run(p.add_run(text), 10, bold)
    return p


def add_evidence(doc, sid: str, caption: str, url: str, width=Inches(6.0)):
    img = ASSETS / f"{sid}.png"
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(cap.add_run(caption), 9, True, RGBColor(0x1F, 0x4E, 0x79))
    if img.exists() and img.stat().st_size > 5000:
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(img), width=width)
    else:
        style_run(
            doc.add_paragraph().add_run(f"[截图 {sid} 未生成]"),
            9, False, RGBColor(0xC0, 0x00, 0x00),
        )
    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(src.add_run(f"来源：{url}"), 8, False, RGBColor(0x66, 0x66, 0x66))
    doc.add_paragraph()


def render_chapter(doc, chapter: dict):
    doc.add_heading(chapter["title"], level=1)
    if chapter.get("intro"):
        add_para(doc, chapter["intro"])
    for sec in chapter.get("sections", []):
        doc.add_heading(sec["heading"], level=2)
        for para in sec.get("paragraphs", []):
            add_para(doc, para)
        if sec.get("table"):
            t = sec["table"]
            add_table(doc, t["headers"], t["rows"])
        if sec.get("note"):
            add_note(doc, sec["note"])


def build_sections_1_to_3(doc):
    doc.add_heading("摘要", level=1)
    add_para(
        doc,
        "本报告聚焦中国及全球 NDR 市场中加密流量检测（ETD）方向，"
        "覆盖 2022–2024 历史数据、2026–2030 中国市场预测、发展趋势与增长机会。"
        "关键数据均附公开来源及原文截图（2026年5月采集）。",
    )
    doc.add_heading("核心发现", level=2)
    add_table(
        doc,
        ["维度", "关键数据", "证据"],
        [
            ["中国 NDR（2024）", "23.6 亿元，-0.9%", "IDC ✅"],
            ["中国 NDR（2030 基准）", "65.7 亿元，CAGR 18.6%", "贝哲斯+【估算】"],
            ["ETD（2030 情景）", "约 23.7 亿元", "【情景分析】"],
            ["加密威胁", "87.2% 经加密通道", "Zscaler ✅"],
            ["HTTPS 普及", "98% 请求", "HTTP Archive ✅"],
        ],
    )

    doc.add_heading("一、中国 NDR 市场规模与格局（2022–2024）", level=1)
    doc.add_heading("1.1 市场规模变化趋势", level=2)
    add_table(
        doc,
        ["年份", "市场规模", "同比增速", "来源"],
        [
            ["2022", "3.5 亿美元（约 24.2 亿元）", "+13.7%", "IDC CHC50358323"],
            ["2023", "23.8 亿元人民币", "+1.6%", "IDC CHC50964624"],
            ["2024", "23.6 亿元人民币", "-0.9%", "IDC CHC52196225"],
        ],
    )
    add_note(doc, "数据冲突：华经 25.48 亿、数世 24.64 亿（2022）— 口径不同，本报告采用 IDC。")
    add_evidence(doc, "S-03", "【图1】IDC 2022：3.5 亿美元", "https://www.secrss.com/articles/58067")
    add_evidence(doc, "S-01", "【图2】IDC 2024：23.6 亿元", "https://www.secrss.com/articles/79613")

    doc.add_heading("1.2 市场竞争格局（2024）", level=2)
    add_table(
        doc,
        ["排名", "厂商", "公开信息"],
        [
            ["1", "奇安信", "连续四年第一"],
            ["2", "绿盟科技", "TOP5"],
            ["3", "深信服科技", "TOP5（与安恒并列第三）"],
            ["4", "安恒信息", "TOP5"],
            ["5", "360 数字安全", "TOP5"],
            ["—", "观成科技", "ETD 代表厂商（IDC 独立章节）"],
        ],
    )
    add_evidence(doc, "S-02", "【图3】IDC 2023：奇安信 22.5%（2023年）", "https://www.qianxin.com/report/detail/rid/62")

    doc.add_heading("二、加密流量检测（ETD）专项分析", level=1)
    doc.add_heading("2.1 全球加密流量威胁态势", level=2)
    add_table(
        doc,
        ["指标", "数值", "来源"],
        [
            ["加密通道攻击占比", "87.2%（+10.3% YoY）", "Zscaler 2024"],
            ["恶意软件占比", "86.5%", "Zscaler 2024"],
            ["制造业攻击占比", "42%", "Zscaler/GlobeNewswire"],
            ["样本量", "321 亿次拦截", "Zscaler 2024"],
        ],
    )
    add_evidence(
        doc, "S-04", "【图4】Zscaler：87.2% 加密威胁",
        "https://www.zscaler.com/blogs/security-research/threatlabz-report-threats-delivered-over-encrypted-channels",
    )

    doc.add_heading("2.2 全球加密流量渗透率", level=2)
    add_table(
        doc,
        ["指标", "数值", "来源"],
        [
            ["HTTPS 请求", "98%", "HTTP Archive 2024"],
            ["HTTPS 站点", "96%", "HTTP Archive 2024"],
            ["TLS 1.3 站点", "73%", "HTTP Archive 2024"],
        ],
    )
    add_evidence(doc, "S-05", "【图5】HTTP Archive：HTTPS 98%", "https://almanac.httparchive.org/en/2024/security")

    doc.add_heading("2.3 IDC 趋势判断", level=2)
    add_para(doc, "IDC 2024：加密流量检测成为更多行业明确需求；AI 不解密检测成为核心竞争力。")
    add_evidence(doc, "S-01", "【图6】IDC 2024 趋势原文", "https://www.secrss.com/articles/79613")

    doc.add_heading("2.4 三大技术路线", level=2)
    add_table(
        doc,
        ["路线", "原理", "代表厂商"],
        [
            ["AI 不解密", "元数据+JA3+行为+AI", "观成、360、绿盟"],
            ["可选解密", "可疑流量选择性解密", "奇安信、深信服"],
            ["全量解密", "中间人解密", "传统厂商"],
        ],
    )

    doc.add_heading("三、全球 NDR 市场对比", level=1)
    add_table(
        doc,
        ["年份", "规模（亿美元）", "CAGR", "来源"],
        [
            ["2024", "33.4", "—", "MarketsandMarkets"],
            ["2025", "37.7", "—", "Grand View"],
            ["2030", "58.2", "9.6%（2025-2030）", "MarketsandMarkets"],
            ["2033", "80.8", "10.1%（2026-2033）", "Grand View"],
        ],
    )
    add_note(doc, "Grand View（2033）与 MarketsandMarkets（2030）不可合并为单一曲线。")
    add_evidence(
        doc, "S-06", "【图7】Grand View：37.7 亿美元",
        "https://www.grandviewresearch.com/industry-analysis/network-detection-response-ndr-market-report",
    )
    add_evidence(
        doc, "S-07", "【图8】MarketsandMarkets：2030 预测",
        "https://www.marketsandmarkets.com/Market-Reports/network-detection-and-response-market-236524642.html",
    )


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    t = doc.add_heading("NDR 市场加密流量方向深度分析报告 V1.2", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "（2022–2024 历史 + 2026–2030 中国预测 + 趋势与机会 · 证据标注版 · 含截图）"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "报告版本：V1.2  |  报告日期：2026年5月  |  NDR 2.0 需求可行性验证")

    build_sections_1_to_3(doc)

    render_chapter(doc, CHAPTER_4)
    add_evidence(
        doc, "S-10", "【图9】数世咨询：NDR 行业需求分布",
        "https://www.dwcon.cn/post/3410",
    )

    render_chapter(doc, CHAPTER_5)
    render_chapter(doc, CHAPTER_6)

    doc.add_heading("七、NDR 2.0 需求验证结论", level=1)
    conclusions = [
        "市场需求成立：87.2% 攻击经加密通道，IDC 将 ETD 列为核心趋势，2030 基准情景下中国 NDR 约 65.7 亿元。",
        "增长路径清晰：金融/运营商/政府为近期主战场，制造/能源为增量重点，云化 ETD 模块可覆盖下沉市场。",
        "技术方向明确：P0 投入 AI 不解密 + TLS/QUIC 元数据；P1 强化 GenAI 降噪与云原生流量。",
        "竞争窗口存在：专精厂商已获 IDC 点名，头部厂商均在布局，需通过 PoC 建立差异化。",
        "建议决策：批准 NDR 2.0 加密流量分析纳入版本规划，Q3 完成竞品实测与 PRD 初稿。",
    ]
    for c in conclusions:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("八、证据截图索引", level=1)
    add_table(
        doc,
        ["图号", "截图", "章节", "URL"],
        [
            ["图1", "S-03", "一、1.1", "secrss.com/articles/58067"],
            ["图2-6", "S-01", "一/二", "secrss.com/articles/79613"],
            ["图3", "S-02", "一、1.2", "qianxin.com/report/detail/rid/62"],
            ["图4", "S-04", "二、2.1", "zscaler.com/..."],
            ["图5", "S-05", "二、2.2", "almanac.httparchive.org/..."],
            ["图7-8", "S-06/S-07", "三", "grandview / marketsandmarkets"],
            ["图9", "S-10", "四、六", "dwcon.cn/post/3410"],
        ],
    )

    doc.add_heading("九、数据来源索引", level=1)
    add_table(
        doc,
        ["#", "机构", "报告/数据", "链接"],
        [
            ["1", "IDC", "NDR 份额 2022-2024", "secrss.com / qianxin.com"],
            ["2", "Zscaler", "Encrypted Attacks 2024", "zscaler.com/blogs/..."],
            ["3", "HTTP Archive", "Web Almanac 2024", "almanac.httparchive.org/..."],
            ["4", "Grand View", "Global NDR Market", "grandviewresearch.com/..."],
            ["5", "MarketsandMarkets", "NDR Forecast 2030", "marketsandmarkets.com/..."],
            ["6", "数世咨询", "NDR 能力指南", "dwcon.cn/post/3410"],
            ["7", "贝哲斯咨询", "中国 NDR 预测", "m.gelonghui.com/p/901401"],
        ],
    )

    add_note(
        doc,
        "免责声明：标注【估算】【情景分析】的数据仅供规划参考；截图以来源网页为准。",
    )

    doc.save(str(OUT))
    doc.save(str(OUT_MAIN))
    print(f"Saved: {OUT}")
    print(f"Saved: {OUT_MAIN}")


if __name__ == "__main__":
    build()
